#!/usr/bin/env python3
"""
create-word-doc.py - 创建符合国家公文标准的 Word 文档

用法:
    ./create-word-doc.py <输出路径> [标题] [内容] [模板类型]

示例:
    ./create-word-doc.py report.docx "项目报告" "## 第一节\n\n内容..." proposal

模板类型:
    proposal(提案) / report(报告) / contract(合同) / minutes(会议纪要) / generic(通用)
"""

import sys
import os
import re
import ssl
import json
import urllib.request
import urllib.error
import xmlrpc.client
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============== 公司信息配置 ==============
USER_HOME = os.path.expanduser("~")
LOGO_DIR = os.path.join(USER_HOME, ".huo15", "assets")
DEFAULT_LOGO_PATH = os.path.join(LOGO_DIR, "logo.png")
FALLBACK_LOGO_URL = 'https://tools.huo15.com/uploads/images/system/logo-colours.png'

# 国标公文页面边距 (GB/T 9704-2012)
MARGIN_TOP = 3.7  # cm
MARGIN_BOTTOM = 3.5  # cm
MARGIN_LEFT = 2.8  # cm
MARGIN_RIGHT = 2.6  # cm


def get_company_info():
    """从公司系统获取公司信息"""
    company_info = {
        'company_name': '青岛火一五信息科技有限公司',
        'vision': '推动B端用户向全场景人工智能机器人转变',
        'philosophy': '打破信息孤岛，用一套系统驱动企业增长',
        'logo_path': None
    }

    # 1. 尝试本地缓存
    if os.path.exists(DEFAULT_LOGO_PATH) and os.path.getsize(DEFAULT_LOGO_PATH) > 1000:
        company_info['logo_path'] = DEFAULT_LOGO_PATH
        return company_info

    # 2. 尝试从公司系统获取
    try:
        agent_id = os.environ.get('OC_AGENT_ID', 'main')
        agents_dir = os.path.expanduser('~/.openclaw/agents')
        creds_file = os.path.join(agents_dir, agent_id, 'odoo_creds.json')

        if os.path.exists(creds_file):
            with open(creds_file, 'r') as f:
                creds = json.load(f)

            openclaw_cfg_file = os.path.expanduser('~/.openclaw/openclaw.json')
            if os.path.exists(openclaw_cfg_file):
                with open(openclaw_cfg_file, 'r') as f:
                    cfg = json.load(f)
                    odoo_env = cfg.get('skills', {}).get('entries', {}).get('huo15-odoo', {}).get('env', {})
                    url = odoo_env.get('ODOO_URL', 'https://huihuoyun.huo15.com')
                    db = odoo_env.get('ODOO_DB', 'huo15_prod')

                user = creds.get('user', '')
                password = creds.get('password', '')

                if user and password:
                    ctx = ssl.create_default_context()
                    ctx.check_hostname = False
                    ctx.verify_mode = ssl.CERT_NONE

                    common = xmlrpc.client.ServerProxy(f'{url}/xmlrpc/2/common', context=ctx)
                    uid = common.authenticate(db, user, password, {})

                    if uid:
                        models = xmlrpc.client.ServerProxy(f'{url}/xmlrpc/2/object', context=ctx)
                        company_data = models.execute_kw(db, uid, password, 'res.company', 'search_read',
                            [[('id', '=', 1)]], {'fields': ['name', 'logo'], 'limit': 1})

                        if company_data:
                            company_info['company_name'] = company_data[0].get('name', company_info['company_name'])
                            logo_id = company_data[0].get('logo')
                            if logo_id:
                                logo_url = f'{url}/web/image/res.company/{logo_id}/logo'
                                ensure_logo_downloaded(logo_url, DEFAULT_LOGO_PATH)
                                if os.path.exists(DEFAULT_LOGO_PATH):
                                    company_info['logo_path'] = DEFAULT_LOGO_PATH
    except Exception as e:
        print(f"获取公司信息失败，使用默认信息: {e}")

    # 3. 备用：下载默认 LOGO
    if not company_info['logo_path']:
        ensure_logo_downloaded(FALLBACK_LOGO_URL, DEFAULT_LOGO_PATH)
        if os.path.exists(DEFAULT_LOGO_PATH):
            company_info['logo_path'] = DEFAULT_LOGO_PATH

    return company_info


def ensure_logo_downloaded(url, dest_path):
    """确保 LOGO 已下载"""
    if os.path.exists(dest_path) and os.path.getsize(dest_path) > 1000:
        return dest_path

    os.makedirs(os.path.dirname(dest_path), exist_ok=True)
    try:
        urllib.request.urlretrieve(url, dest_path)
        print(f"✓ LOGO 已下载: {dest_path}")
    except Exception as e:
        print(f"⚠ LOGO 下载失败: {e}")
        try:
            urllib.request.urlretrieve(FALLBACK_LOGO_URL, dest_path)
        except:
            pass
    return dest_path


def set_chinese_font(run, font_name='仿宋', font_size=12, bold=False):
    """设置中文字体，确保 WPS 和 Word 兼容"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    return run


def add_header_with_logo(doc, logo_path=None, company_name=None):
    """添加页眉：LOGO + 公司名称（靠左对齐，紧密挨着）"""
    if logo_path is None or company_name is None:
        info = get_company_info()
        logo_path = logo_path or info.get('logo_path')
        company_name = company_name or info.get('company_name', '青岛火一五信息科技有限公司')

    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # 清除默认段落
    for para in header.paragraphs:
        para.clear()

    # 创建第一个段落：LOGO + 公司名称（同一行，靠左）
    if not header.paragraphs:
        para = header.add_paragraph()
    else:
        para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 添加 LOGO（高度 1.0cm）
    if logo_path and os.path.exists(logo_path):
        try:
            run = para.add_run()
            run.add_picture(logo_path, height=Cm(1.0))
        except Exception as e:
            print(f"⚠ LOGO 添加失败: {e}")

    # 添加公司名称（紧跟 LOGO，中间一个空格）
    run = para.add_run(f' {company_name}')
    set_chinese_font(run, '黑体', 10)

    # 添加底线（分隔线）
    pPr = OxmlElement('w:pPr')
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    para._element.insert(0, pPr)

    return header


def add_footer_with_page_numbers(doc):
    """添加页脚：第 X 页 共 Y 页（居中，仿宋小四）"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    if not footer.paragraphs:
        para = footer.add_paragraph()
    else:
        para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 清除内容
    for run in para.runs:
        run.text = ''

    # "第 "
    run1 = para.add_run('第 ')
    set_chinese_font(run1, '仿宋', 10.5)

    # PAGE 域
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText1 = OxmlElement('w:instrText')
    instrText1.set(qn('xml:space'), 'preserve')
    instrText1.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run_page = para.add_run()
    run_page._element.clear()
    run_page._element.append(fldChar1)
    run_page._element.append(instrText1)
    run_page._element.append(fldChar2)
    set_chinese_font(run_page, '仿宋', 10.5)

    # " 页 共 "
    run2 = para.add_run(' 页 共 ')
    set_chinese_font(run2, '仿宋', 10.5)

    # NUMPAGES 域
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = ' NUMPAGES '
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    run_total = para.add_run()
    run_total._element.clear()
    run_total._element.append(fldChar3)
    run_total._element.append(instrText2)
    run_total._element.append(fldChar4)
    set_chinese_font(run_total, '仿宋', 10.5)

    # " 页"
    run3 = para.add_run(' 页')
    set_chinese_font(run3, '仿宋', 10.5)


def parse_markdown_to_word(doc, content):
    """将 markdown 内容转换为 Word 格式（不保留 markdown 语法）"""

    lines = content.split('\n')
    i = 0
    table_buffer = []

    def flush_table():
        """将缓存的表格行输出为 Word 表格"""
        nonlocal table_buffer
        if not table_buffer:
            return
        # 解析表头和行
        rows_data = []
        for line in table_buffer:
            # 跳过分割线
            if re.match(r'^[\|\-\s]+$', line):
                continue
            cells = [c.strip() for c in line.strip('|').split('|')]
            rows_data.append(cells)

        if len(rows_data) < 2:
            table_buffer = []
            return

        # 创建表格
        table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
        table.style = 'Table Grid'

        for r, row in enumerate(rows_data):
            for c, cell_text in enumerate(row):
                cell = table.rows[r].cells[c]
                cell.text = cell_text
                # 设置单元格字体
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if r == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in para.runs:
                        if r == 0:
                            set_chinese_font(run, '黑体', 10.5, bold=True)
                        else:
                            set_chinese_font(run, '仿宋', 10.5)

        table_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            flush_table()
            doc.add_paragraph()
            i += 1
            continue

        # 跳过 markdown 代码块标记
        if stripped.startswith('```'):
            i += 1
            continue

        # 表格行（以 | 开头）
        if stripped.startswith('|'):
            # 检查是否是分割线
            if re.match(r'^[\|\-\s]+$', stripped):
                i += 1
                continue
            table_buffer.append(stripped)
            i += 1
            continue
        else:
            flush_table()

        # 标题处理（去除 # 符号，转为正规标题样式）
        if stripped.startswith('#### '):
            text = stripped[5:]
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_chinese_font(run, '仿宋', 12, bold=True)
            p.paragraph_format.line_spacing = 1.5
        elif stripped.startswith('### '):
            text = stripped[4:]
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_chinese_font(run, '仿宋', 12, bold=True)
            p.paragraph_format.line_spacing = 1.5
        elif stripped.startswith('## '):
            text = stripped[3:]
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_chinese_font(run, '楷体', 15, bold=True)
            p.paragraph_format.line_spacing = 1.5
        elif stripped.startswith('# '):
            text = stripped[2:]
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_chinese_font(run, '黑体', 15, bold=True)
            p.paragraph_format.line_spacing = 1.5
        # 列表处理
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            text = parse_inline_formatting(text)
            p = doc.add_paragraph(text, style='List Bullet')
            for run in p.runs:
                set_chinese_font(run, '仿宋', 12)
            p.paragraph_format.line_spacing = 1.5
        elif re.match(r'^\d+\.\s+', stripped):
            # 有序列表
            text = re.sub(r'^\d+\.\s+', '', stripped)
            text = parse_inline_formatting(text)
            p = doc.add_paragraph(text, style='List Number')
            for run in p.runs:
                set_chinese_font(run, '仿宋', 12)
            p.paragraph_format.line_spacing = 1.5
        else:
            # 普通段落，处理行内格式（**bold** 等）
            text = parse_inline_formatting(stripped)
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_chinese_font(run, '仿宋', 12)
            p.paragraph_format.line_spacing = 1.5

        i += 1

    # 处理最后残留的表格
    flush_table()


def parse_inline_formatting(text):
    """解析行内格式：**bold**、*italic*、[text](url)"""
    # 去除 markdown 链接 [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    # 去除 markdown 图片 ![alt](url) -> alt
    text = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', r'\1', text)
    return text


def create_word_doc(output_path, title="", content="", template="generic"):
    """
    创建符合国家公文标准的 Word 文档

    参数:
        output_path: 输出文件路径
        title: 文档标题
        content: 文档内容（支持 markdown 格式）
        template: 模板类型
    """

    # 1. 创建文档
    doc = Document()

    # 2. 设置页面边距（GB/T 9704-2012）
    for section in doc.sections:
        section.top_margin = Cm(MARGIN_TOP)
        section.bottom_margin = Cm(MARGIN_BOTTOM)
        section.left_margin = Cm(MARGIN_LEFT)
        section.right_margin = Cm(MARGIN_RIGHT)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)

    # 3. 添加页眉（LOGO + 公司名称，靠左对齐）
    add_header_with_logo(doc)

    # 4. 添加页脚（页码）
    add_footer_with_page_numbers(doc)

    # 5. 设置默认样式
    style = doc.styles['Normal']
    style.font.name = '仿宋'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    # 6. 添加标题（居中，黑体二号加粗）
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        set_chinese_font(run, '黑体', 22, bold=True)
        p.paragraph_format.line_spacing = 1.5

    # 7. 解析 markdown 内容
    if content:
        parse_markdown_to_word(doc, content)

    # 8. 保存
    doc.save(output_path)
    print(f"✅ Word 文档已生成: {output_path}")


if __name__ == "__main__":
    output = sys.argv[1] if len(sys.argv) > 1 else "output.docx"
    title = sys.argv[2] if len(sys.argv) > 2 else ""
    content = sys.argv[3] if len(sys.argv) > 3 else ""
    template = sys.argv[4] if len(sys.argv) > 4 else "generic"

    create_word_doc(output, title, content, template)
