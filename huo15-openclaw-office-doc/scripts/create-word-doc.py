#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
create-word-doc.py — 火一五企业级 Word 文档生成器 v7.0

相对 v6.0 的改进：
  - 解析器拆出 doc_core.py，与 PDF 渲染器共用
  - 修复 CJK 软换行（不再插入多余空格）+ 支持硬换行（行尾 2 空格 / `\\`）
  - 强制页眉左对齐：直接写 OOXML jc，并清除 Header 样式自带的 tab stops
  - 新增 6 类规范：商业计划书 / 用户手册 / 培训手册 / 招投标书 / 演讲稿 / 研究报告
  - 列表 / 引用 / 元数据 / 表头 单元格也支持硬换行（同 paragraph 路径）
  - render_inline 内的硬换行通过 `\\n` 转 <w:br/>
  - 显式分页符：`---PAGE---` / `\\pagebreak`

用法：
    python create-word-doc.py --output 文档.docx --title '标题' --content '...'
兼容旧位置参数：
    python create-word-doc.py <输出路径> [标题] [正文] [编号] [版本] [密级] [格式]
"""

import sys
import os
import json
import argparse
import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 共享解析器 + 规范预设
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import doc_core  # noqa: E402

# Pygments 是可选依赖：装了走代码高亮，没装回落纯文本灰色
try:
    from pygments import lex
    from pygments.lexers import get_lexer_by_name
    from pygments.util import ClassNotFound
    HAS_PYGMENTS = True
except ImportError:  # pragma: no cover
    HAS_PYGMENTS = False


# ============================================================
# 一、OOXML 小工具
# ============================================================

def _set_font(run, font_name, size, bold=False, italic=False, color=None):
    """统一设置中英文字体（WPS / Word 双兼容）。"""
    run.font.name = font_name
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _force_paragraph_alignment(paragraph, ooxml_val='left',
                               clear_tabs=False):
    """直接写 OOXML 的 jc 元素。比 python-docx 的 alignment 属性更可靠：

    - python-docx 在某些 style 继承场景下不会真的 emit `<w:jc>`
    - WPS 在 Header 样式上会无视 alignment 属性，必须显式 jc
    可选清除 Header 样式自带的 tab stops（中部 / 右部）以避免视觉错位。
    """
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        paragraph._element.insert(0, pPr)

    for jc in pPr.findall(qn('w:jc')):
        pPr.remove(jc)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), ooxml_val)
    pPr.append(jc)

    if clear_tabs:
        for tabs in pPr.findall(qn('w:tabs')):
            pPr.remove(tabs)


def _add_border_bottom(paragraph, color='888888', sz='6'):
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        paragraph._element.insert(0, pPr)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_border_left(paragraph, color='CCCCCC', sz='18'):
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        paragraph._element.insert(0, pPr)
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)


def _set_cell_shading(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)


def _add_cjk_paragraph_props(paragraph):
    """为段落开启 Word 内置的 CJK 排版优化：
       - autoSpaceDE：英文与中文间自动加空格
       - autoSpaceDN：数字与中文间自动加空格
       - kinsoku（标点挤压）：避免行首禁则字（。」），行末禁则字（「）出现
       - wordWrap：允许英文单词在行尾被打断（提升中文段落对齐效果）
    这几个属性是 Word/WPS 的标准段落属性，不需要新依赖。
    """
    pPr = paragraph._element.get_or_add_pPr()
    for tag, val in [
        ('w:autoSpaceDE', '1'),
        ('w:autoSpaceDN', '1'),
        ('w:kinsoku', '1'),
        ('w:wordWrap', '1'),
        ('w:overflowPunct', '1'),
        ('w:topLinePunct', '0'),
    ]:
        for el in pPr.findall(qn(tag)):
            pPr.remove(el)
        el = OxmlElement(tag)
        el.set(qn('w:val'), val)
        pPr.append(el)


def _set_first_line_indent_chars(paragraph, chars=2):
    """以"字符数"设置首行缩进（200 = 2 字符），跨字号保持视觉一致。

    比起物理 cm 值（python-docx 默认），firstLineChars 更符合中文公文规范——
    标准要求"首行缩进 2 字符"，而 cm 在不同字号下视觉效果就会跑偏。
    """
    pPr = paragraph._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    # 清除任何已存在的 firstLine（cm dxa 单位），避免双重缩进
    for attr in ['firstLine']:
        if ind.get(qn(f'w:{attr}')) is not None:
            del ind.attrib[qn(f'w:{attr}')]
    ind.set(qn('w:firstLineChars'), str(int(chars * 100)))


def _set_paragraph_shading(paragraph, fill_color):
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        paragraph._element.insert(0, pPr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    pPr.append(shd)


def _add_field(paragraph, field_code, font_name, font_size):
    run = paragraph.add_run()
    fc_begin = OxmlElement('w:fldChar')
    fc_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' {field_code} '
    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    run._element.append(fc_begin)
    run._element.append(instr)
    run._element.append(fc_end)
    _set_font(run, font_name, font_size)


# ============================================================
# 二、页眉 & 页脚
# ============================================================

def _reset_header_paragraph(para, alignment='left'):
    """彻底清空一个 header / footer 中的段落，并强制对齐。

    不只清 run 文本——把段落已有 runs 全部移除，避免 WPS 里残留旧 LOGO。
    然后用 OOXML 直接写 jc，并清掉 Header 样式继承的 tab stops。
    """
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.left_indent = Cm(0)
    para.paragraph_format.right_indent = Cm(0)
    para.paragraph_format.first_line_indent = Cm(0)
    _force_paragraph_alignment(para, ooxml_val=alignment, clear_tabs=True)


def build_header(doc, preset, logo_path, company_name, doc_number=None,
                 classification=None, title=''):
    """页眉：
       - 默认（公文 / 方案 / 需求 / 招投 / 工作报告等）→ LOGO + 名称（左对齐）
                                                       + 文档编号 + 密级
       - 'centered'（合同）→ 仅公司名居中
       - 'minimal'（用户手册 / 演讲稿）→ LOGO + 名称（左对齐），不显示编号 / 密级
    """
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    if preset.header_layout == 'centered':
        _reset_header_paragraph(para, alignment='center')
    else:
        _reset_header_paragraph(para, alignment='left')

    # LOGO（centered 也带）
    if logo_path and os.path.exists(logo_path):
        try:
            run = para.add_run()
            run.add_picture(logo_path, height=Cm(0.9))
            sep = para.add_run('  ')
            _set_font(sep, '黑体', preset.size_body - 2)
        except Exception as exc:  # pragma: no cover
            print(f'⚠️  LOGO 添加失败：{exc}', file=sys.stderr)

    run = para.add_run(company_name)
    _set_font(run, '黑体', preset.size_body - 2)

    # 编号 / 密级 仅在 'company' 布局且非 minimal 时显示
    if preset.header_layout == 'company':
        if doc_number:
            run = para.add_run(f'    {doc_number}')
            _set_font(run, '黑体', preset.size_body - 2)
        if classification:
            run = para.add_run(f'    【{classification}】')
            _set_font(run, '黑体', preset.size_body - 2, bold=True)

    _add_border_bottom(para, color='888888', sz='6')


def build_footer(doc, preset, company_name):
    """页脚：所有规范统一为 第 X 页 / 共 Y 页（PAGE / NUMPAGES 字段）。"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    _reset_header_paragraph(para, alignment='center')

    size = preset.size_body - 2

    run = para.add_run('第 ')
    _set_font(run, preset.font_body, size)
    _add_field(para, 'PAGE', preset.font_body, size)
    run = para.add_run(' 页 / 共 ')
    _set_font(run, preset.font_body, size)
    _add_field(para, 'NUMPAGES', preset.font_body, size)
    run = para.add_run(' 页')
    _set_font(run, preset.font_body, size)


# ============================================================
# 三、Inline 渲染（包含硬换行）
# ============================================================

def render_inline(paragraph, text, font, size, base_bold=False,
                  color=None, inline_code_font='Consolas'):
    """把一段（含 `\\n` 硬换行）写到 paragraph，按 Block 内联标记拆分 run。"""
    if not text:
        return

    # text 中的 '\n' 表示硬换行：先按行拆，每行内再拆内联 token
    for line_idx, line in enumerate(doc_core.split_paragraph_lines(text)):
        if line_idx > 0:
            br_run = paragraph.add_run()
            _set_font(br_run, font, size)
            br_run.add_break(WD_BREAK.LINE)

        for kind, payload in doc_core.tokenize_inline(line):
            if kind == 'bold':
                run = paragraph.add_run(payload)
                _set_font(run, font, size, bold=True, color=color)
            elif kind == 'italic':
                run = paragraph.add_run(payload)
                _set_font(run, font, size, bold=base_bold,
                          italic=True, color=color)
            elif kind == 'code':
                run = paragraph.add_run(payload)
                _set_font(run, inline_code_font, size, bold=base_bold)
            else:
                run = paragraph.add_run(payload)
                _set_font(run, font, size, bold=base_bold, color=color)


def _apply_paragraph_defaults(p, preset, indent=True, align=None,
                              space_before=0, space_after=None):
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = preset.line_spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(
        space_after if space_after is not None else preset.paragraph_spacing_pt
    )
    # 首行缩进：preset 缩进 > 0 时按字符化处理（公文规范），否则清零
    if indent and preset.first_line_indent_cm > 0:
        _set_first_line_indent_chars(p, chars=2)
    else:
        p.paragraph_format.first_line_indent = Cm(0)
    # 中英自动空格 + 标点挤压（所有正文段落统一应用）
    _add_cjk_paragraph_props(p)


# ============================================================
# 四、Block 渲染
# ============================================================

# ----- 书签 / TOC 字段支持 -----

_BOOKMARK_COUNTER = [0]


def _reset_bookmarks():
    _BOOKMARK_COUNTER[0] = 0


def _add_heading_bookmark(paragraph, level):
    """给标题段落两端插入 _Toc 书签，TOC 字段才能识别这条标题。"""
    _BOOKMARK_COUNTER[0] += 1
    bk_id = _BOOKMARK_COUNTER[0]
    bk_name = f'_Toc{str(bk_id).zfill(8)}'

    bk_start = OxmlElement('w:bookmarkStart')
    bk_start.set(qn('w:id'), str(bk_id))
    bk_start.set(qn('w:name'), bk_name)

    bk_end = OxmlElement('w:bookmarkEnd')
    bk_end.set(qn('w:id'), str(bk_id))

    p_el = paragraph._element
    pPr = p_el.find(qn('w:pPr'))
    insert_after = pPr if pPr is not None else None
    if insert_after is not None:
        insert_after.addnext(bk_start)
    else:
        p_el.insert(0, bk_start)
    p_el.append(bk_end)
    return bk_name


def add_toc_field(doc, preset, levels='1-3', title='目录'):
    """v7.6.0：插入 TOC 字段并返回 (toc_paragraph, end_run) 引用，
    供 _finalize_static_toc 在 render 完后回填收集到的标题。

    旧版用一个永久占位符 "目录将在打开 Word / WPS 时自动生成"，看起来像正文。
    新版先用同一占位符，但在所有 heading 渲染完后调 _finalize_static_toc，
    把占位符替换成 H1-H3 标题列表（带级别缩进 + `<w:br>` 换行）。
    Word/WPS 用户打开 → updateFields=true 触发 F9 → 替换为带页码的真目录。
    """
    head_p = doc.add_paragraph()
    _apply_paragraph_defaults(head_p, preset, indent=False,
                              align=WD_ALIGN_PARAGRAPH.LEFT,
                              space_before=12, space_after=8)
    head_run = head_p.add_run(title)
    _set_font(head_run, preset.font_heading, preset.size_section, bold=True)

    p = doc.add_paragraph()
    _apply_paragraph_defaults(p, preset, indent=False,
                              align=WD_ALIGN_PARAGRAPH.LEFT,
                              space_before=0, space_after=4)
    # 字段三段式：begin / instr / separate / 占位 / end
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    fld_begin.set(qn('w:dirty'), 'true')
    run._element.append(fld_begin)

    run = p.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' TOC \\o "{levels}" \\h \\z \\u '
    run._element.append(instr)

    run = p.add_run()
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run._element.append(fld_sep)

    # 临时占位（render 后会被 _finalize_static_toc 用 collected headings 替换）
    placeholder_run = p.add_run()
    _set_font(placeholder_run, preset.font_body, preset.size_body)

    end_run = p.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    end_run._element.append(fld_end)

    # 分页
    pp = doc.add_paragraph()
    pp.add_run().add_break(WD_BREAK.PAGE)

    return p, placeholder_run


# v7.6.0: 渲染期间收集的 (level, text) — render_heading 写入；
# _finalize_static_toc 读出后清空。
_TOC_HEADINGS = []


def _reset_toc_collector():
    _TOC_HEADINGS.clear()


def _collect_heading(level, text):
    if 1 <= level <= 3:
        _TOC_HEADINGS.append((level, text))


def _finalize_static_toc(toc_paragraph, placeholder_run, preset):
    """v7.6.0：把 _TOC_HEADINGS 里的 (level, text) 写到 TOC 字段缓存里，
    取代占位符。打开 Word / WPS 前用户就能看到完整目录列表（无页码）；
    用户打开后 updateFields=true 触发字段刷新，替换为带真页码的 TOC。
    """
    if toc_paragraph is None or placeholder_run is None:
        return
    headings = list(_TOC_HEADINGS)
    if not headings:
        # 没有任何 H1-H3 → 把整段 TOC 段落连同标题段落一起隐去（设极小空白）
        # 简单做法：占位符填一个空白字符，让段落看起来空但保持字段结构
        placeholder_run.text = ''
        return

    # 准备 H1/H2/H3 各自的字号 / 缩进
    body_size = preset.size_body
    placeholder_run.text = ''
    parent = placeholder_run._element.getparent()
    placeholder_idx = list(parent).index(placeholder_run._element)

    new_runs = []
    for i, (level, text) in enumerate(headings):
        # 缩进：H1 无缩进，H2 两个全角空格，H3 四个全角空格
        indent = '　　' * max(level - 1, 0)
        line = f'{indent}{text}'
        run_el = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        # 字体
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), preset.font_body)
        rFonts.set(qn('w:hAnsi'), preset.font_body)
        rFonts.set(qn('w:eastAsia'), preset.font_body)
        rPr.append(rFonts)
        # 字号（half-points）；H1 略大
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str((body_size + (1 if level == 1 else 0)) * 2))
        rPr.append(sz)
        # 颜色：H1 黑，H2 中灰，H3 浅灰（视觉层次）
        col = OxmlElement('w:color')
        col.set(qn('w:val'),
                {1: '222222', 2: '555555', 3: '888888'}.get(level, '555555'))
        rPr.append(col)
        run_el.append(rPr)
        # 文本
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = line
        run_el.append(t)
        new_runs.append(run_el)

        # 行间 break（最后一行不加）
        if i < len(headings) - 1:
            br_run = OxmlElement('w:r')
            br = OxmlElement('w:br')
            br_run.append(br)
            new_runs.append(br_run)

    # 倒序插入到 placeholder 位置（保留 fld_end 在最后）
    for new_el in reversed(new_runs):
        parent.insert(placeholder_idx, new_el)


def _set_update_fields_on_open(doc):
    """settings.xml 里加 updateFields=true，让 Word/WPS 打开时自动刷字段。"""
    settings = doc.settings.element
    for el in settings.findall(qn('w:updateFields')):
        settings.remove(el)
    update = OxmlElement('w:updateFields')
    update.set(qn('w:val'), 'true')
    settings.append(update)


def set_doc_core_properties(doc, title=None, author=None, subject=None,
                             keywords=None, category=None, comments=None):
    """填 docx 的 core properties。便于 OA 检索 / 投标书系统索引。"""
    cp = doc.core_properties
    if title:
        cp.title = title[:255]
    if author:
        cp.author = author
        cp.last_modified_by = author
    if subject:
        cp.subject = subject
    if keywords:
        cp.keywords = keywords if isinstance(keywords, str) else ','.join(keywords)
    if category:
        cp.category = category
    if comments:
        cp.comments = comments
    now = datetime.datetime.now()
    cp.created = now
    cp.modified = now


# ----- Pygments 代码块高亮 -----

def _pygments_color(token_type):
    """token_type 是 pygments.token.Token 对象。返回 RGBColor 或 None。"""
    s = repr(token_type)
    hexcol = doc_core.get_token_color(s)
    if hexcol:
        return RGBColor(int(hexcol[0:2], 16), int(hexcol[2:4], 16),
                        int(hexcol[4:6], 16))
    return None


def render_heading(doc, preset, level, text):
    _collect_heading(level, text)

    if level <= 1:
        font_size, bold = preset.size_chapter, True
        space_before, space_after = 18, 8
    elif level == 2:
        font_size, bold = preset.size_section, True
        space_before, space_after = 14, 6
    else:
        font_size, bold = preset.size_body + 1, True
        space_before, space_after = 8, 4

    p = doc.add_paragraph()
    _apply_paragraph_defaults(p, preset, indent=False,
                              align=WD_ALIGN_PARAGRAPH.LEFT,
                              space_before=space_before,
                              space_after=space_after)

    # 给 Word 内置导航识别：手动写 outlineLvl（前 6 级映射 0~5）
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p._element.insert(0, pPr)
    for el in pPr.findall(qn('w:outlineLvl')):
        pPr.remove(el)
    outline = OxmlElement('w:outlineLvl')
    outline.set(qn('w:val'), str(min(level, 6) - 1))
    pPr.append(outline)

    # 加 _Toc 书签，TOC 字段才能识别
    if level <= 3:
        _add_heading_bookmark(p, level)

    render_inline(p, text, preset.font_heading, font_size, base_bold=bold)


def render_paragraph(doc, preset, text):
    p = doc.add_paragraph()
    _apply_paragraph_defaults(p, preset, indent=True)
    render_inline(p, text, preset.font_body, preset.size_body)


def render_list(doc, preset, ordered, items):
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        _apply_paragraph_defaults(p, preset, indent=False,
                                  align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                  space_after=3)
        bullet = f'{idx}. ' if ordered else '• '
        run = p.add_run(bullet)
        _set_font(run, preset.font_body, preset.size_body)
        render_inline(p, item, preset.font_body, preset.size_body)


def render_table(doc, preset, rows, has_header=True):
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    norm = [r + [''] * (max_cols - len(r)) for r in rows]
    table = doc.add_table(rows=len(norm), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row in enumerate(norm):
        is_head = has_header and r_idx == 0
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ''
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            font_name = preset.font_heading if is_head else preset.font_body
            size = preset.size_body - 1
            render_inline(para, cell_text, font_name, size,
                          base_bold=is_head)
            if is_head:
                _set_cell_shading(cell, 'E8ECF0')


def _emit_code_run(para, text, color, size):
    """代码 run；'\n' 转换成软回车，单 run 内不留 break，颜色统一。"""
    if not text:
        return
    parts = text.split('\n')
    for idx, part in enumerate(parts):
        if idx > 0:
            br_run = para.add_run()
            _set_font(br_run, 'Consolas', size)
            br_run.add_break(WD_BREAK.LINE)
        if part:
            run = para.add_run(part)
            _set_font(run, 'Consolas', size,
                      color=color or RGBColor(0x22, 0x22, 0x22))


def render_code_block(doc, preset, code, lang=''):
    """带浅灰背景 + 等宽字体的代码段；装了 Pygments 走 token 着色。"""
    para = doc.add_paragraph()
    _apply_paragraph_defaults(para, preset, indent=False,
                              align=WD_ALIGN_PARAGRAPH.LEFT,
                              space_before=4, space_after=8)
    para.paragraph_format.line_spacing = 1.25
    para.paragraph_format.left_indent = Cm(0.3)
    _set_paragraph_shading(para, 'F7F7F7')
    _add_border_bottom(para, color='CCCCCC', sz='4')

    size = preset.size_body - 1

    if lang:
        tag = para.add_run(f'{lang}\n')
        _set_font(tag, 'Consolas', preset.size_body - 2,
                  color=RGBColor(0x88, 0x88, 0x88))

    # 1) 尝试 Pygments 高亮
    if HAS_PYGMENTS and lang:
        try:
            lexer = get_lexer_by_name(lang, stripall=False, ensurenl=False)
            for ttype, value in lex(code, lexer):
                _emit_code_run(para, value, _pygments_color(ttype), size)
            return
        except (ClassNotFound, Exception):
            pass  # 回落到无颜色

    # 2) 无颜色：保留原有等宽 + 灰色文本
    _emit_code_run(para, code, RGBColor(0x22, 0x22, 0x22), size)


def render_blockquote(doc, preset, lines):
    for line in lines:
        p = doc.add_paragraph()
        _apply_paragraph_defaults(p, preset, indent=False,
                                  align=WD_ALIGN_PARAGRAPH.LEFT,
                                  space_after=2)
        p.paragraph_format.left_indent = Cm(0.5)
        _add_border_left(p, color='FF7043', sz='18')
        render_inline(p, line, preset.font_body, preset.size_body,
                      color=RGBColor(0x55, 0x55, 0x55))


def render_metadata(doc, preset, pairs):
    if not pairs:
        return
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r, (key, value) in enumerate(pairs):
        row = table.rows[r]
        row.cells[0].text = ''
        row.cells[1].text = ''
        p0 = row.cells[0].paragraphs[0]
        p1 = row.cells[1].paragraphs[0]
        render_inline(p0, key or '', preset.font_heading,
                      preset.size_body - 1, base_bold=True)
        render_inline(p1, value or '', preset.font_body,
                      preset.size_body - 1)
        _set_cell_shading(row.cells[0], 'F5F5F5')


def render_hr(doc, preset):
    p = doc.add_paragraph()
    _apply_paragraph_defaults(p, preset, indent=False,
                              space_before=4, space_after=6)
    _add_border_bottom(p, color='CCCCCC', sz='6')


def render_page_break(doc, preset):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def render_block(doc, preset, block):
    btype = block['type']
    if btype == 'heading':
        render_heading(doc, preset, block['level'], block['text'])
    elif btype == 'paragraph':
        text = block['text']
        detected = doc_core.detect_heading_from_preset(text, preset)
        if detected is not None:
            level, cleaned = detected
            render_heading(doc, preset, level, cleaned)
        else:
            render_paragraph(doc, preset, text)
    elif btype == 'list':
        render_list(doc, preset, block.get('ordered', False),
                    block.get('items', []))
    elif btype == 'table':
        render_table(doc, preset, block.get('rows', []),
                     has_header=block.get('has_header', True))
    elif btype == 'code_block':
        render_code_block(doc, preset, block.get('code', ''),
                          block.get('lang', ''))
    elif btype == 'blockquote':
        render_blockquote(doc, preset, block.get('lines', []))
    elif btype == 'metadata':
        render_metadata(doc, preset, block.get('pairs', []))
    elif btype == 'hr':
        render_hr(doc, preset)
    elif btype == 'page_break':
        render_page_break(doc, preset)


def render_content(doc, preset, content):
    blocks = doc_core.parse_blocks(content)
    if not blocks:
        p = doc.add_paragraph()
        _apply_paragraph_defaults(p, preset, indent=True, space_after=0)
        run = p.add_run('（无正文内容）')
        _set_font(run, preset.font_body, preset.size_body,
                  color=RGBColor(0x99, 0x99, 0x99))
        return
    for block in blocks:
        render_block(doc, preset, block)


# ============================================================
# 五、文档壳（标题、元数据、版本历史、审批表）
# ============================================================

_AUTO_META_KEYS = frozenset({
    '文档编号', '编号', '合同编号', '协议编号', '订单编号', '报价编号',
    '版本', '版次', '密级', '机密等级',
    '日期', '签订日期', '签约日期', '验收日期', '生效日期',
    '作者', '编制', '审核', '批准',
})


def _looks_like_meta_kv(pairs):
    """≥1 对命中 _AUTO_META_KEYS 即视为元数据集合。"""
    return any((k or '').strip().strip('*~`') in _AUTO_META_KEYS
               for k, _ in (pairs or []))


def _looks_like_meta_table(rows):
    """v7.6.0 检测 markdown TABLE 形式的元数据：
       - 2 列（或多列但实际数据 2 列）
       - 第一列至少 2 个 cell 命中 _AUTO_META_KEYS
    """
    if not rows or len(rows) < 2:
        return False
    if not all(len(r) >= 2 for r in rows):
        return False
    first_col = [(r[0] or '').strip().strip('*~`') for r in rows]
    hits = sum(1 for k in first_col if k in _AUTO_META_KEYS)
    return hits >= 2


def _heading_count_ge(content, threshold):
    """v7.6.0：统计正文里的 H1+H2 数量，便于判断 TOC 是否值得生成。"""
    if not content:
        return False
    count = 0
    for b in doc_core.parse_blocks(content):
        if b.get('type') == 'heading' and b.get('level', 9) <= 2:
            count += 1
            if count >= threshold:
                return True
    return False


def _body_has_top_metadata_block(content):
    """v7.6.0：检查 markdown 正文最前面（跳过空行 + 可选 H1）是否已有元数据。

    覆盖 LLM 生成元数据的两种常见写法：

    (1) KV style（被 parse_blocks 归并为 'metadata' block）：
        **文档编号：** HG-TD-2026-001
        **版本：** V1.0
        **密级：** 内部

    (2) TABLE style（普通 markdown 表）：
        | 文档编号 | HG-TD-2026-001 |
        | 版本    | V1.0           |
        | 密级    | 内部           |

    命中后跳过 add_doc_meta，避免和 CLI 自动元数据表重复。
    """
    if not content:
        return False
    blocks = doc_core.parse_blocks(content)
    if not blocks:
        return False
    first = blocks[0]
    # 跳过最多一个 H1
    if first.get('type') == 'heading' and first.get('level') == 1:
        if len(blocks) < 2:
            return False
        first = blocks[1]
    btype = first.get('type')
    if btype == 'metadata':
        return _looks_like_meta_kv(first.get('pairs'))
    if btype == 'table':
        return _looks_like_meta_table(first.get('rows'))
    return False


def _maybe_dedupe_h1_title(content, title, preset, want_title_block):
    """若 markdown 首个非空块是 H1 且文本等于 --title，剥掉这一行。

    v7.3：常见模式是 LLM 同时在 --title 和正文 `# Title` 都写一份；旧版渲染
    出来标题会重复出现。本函数只剥掉"等同于 --title 的首个 H1"，对正文里
    后续 `# 别的章节` 不动。
    """
    if not content or not title:
        return content
    if not want_title_block:
        return content
    if not getattr(preset, 'dedupe_h1_title', True):
        return content

    target = _strip_markdown_emphasis(title.strip())
    if not target:
        return content

    lines = content.split('\n')
    # 跳过 Markdown 首部空行和 yaml frontmatter（--- ... ---）
    i = 0
    n = len(lines)
    while i < n and not lines[i].strip():
        i += 1
    if i < n and lines[i].strip() == '---':
        # YAML frontmatter（如果有）— 找下一个 ---
        j = i + 1
        while j < n and lines[j].strip() != '---':
            j += 1
        if j < n:
            i = j + 1
            while i < n and not lines[i].strip():
                i += 1

    if i >= n:
        return content
    head = lines[i].strip()
    if not head.startswith('# '):
        return content
    head_text = _strip_markdown_emphasis(head[2:].strip())
    if head_text != target:
        return content

    new_lines = lines[:i] + lines[i + 1:]
    return '\n'.join(new_lines)


def _strip_markdown_emphasis(text):
    """剥掉首尾的 markdown 强调标记 (`**` / `*` / ``` ` ```)。

    v7.3：标题和文档元数据这种"理应是纯文本"的字段经常被 LLM 顺手包成
    `**Foo**`，而旧 add_title / add_doc_meta 直接 add_run，导致 `**` 留在输出里。
    用这个工具先把外层标记剥掉，再交给 render_inline 处理可能仍存在的内部
    markdown。
    """
    if not text:
        return text
    s = text.strip()
    while True:
        changed = False
        for marker in ('**', '*', '`', '~~'):
            if (len(s) >= 2 * len(marker)
                    and s.startswith(marker) and s.endswith(marker)):
                s = s[len(marker):-len(marker)].strip()
                changed = True
                break
        if not changed:
            break
    return s


def add_title(doc, preset, title):
    if not title:
        return
    title = _strip_markdown_emphasis(title)
    if not title:
        return
    p = doc.add_paragraph()
    if getattr(preset, 'title_alignment', 'center') == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = preset.line_spacing
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    # 仍然走 render_inline，让用户偶尔在 title 里夹的 `*` / `` ` `` 等内部
    # markdown 也能正确呈现；外层 ** 已被 _strip_markdown_emphasis 剥掉。
    render_inline(p, title, preset.font_title, preset.size_title,
                  base_bold=True)


def add_doc_meta(doc, preset, doc_number, version, classification, author):
    # v7.3 防御：CLI 传进来的字段偶尔被 LLM 包成 `**X**`，剥一下
    doc_number = _strip_markdown_emphasis(doc_number) if doc_number else doc_number
    version = _strip_markdown_emphasis(version) if version else version
    classification = (_strip_markdown_emphasis(classification)
                      if classification else classification)
    author = _strip_markdown_emphasis(author) if author else author

    items = []
    if doc_number:
        items.append(('文档编号', doc_number))
    items.append(('版本', version))
    items.append(('密级', classification))
    items.append(('日期', datetime.date.today().strftime('%Y-%m-%d')))
    if author:
        items.append(('作者', author))

    table = doc.add_table(rows=len(items), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r, (key, value) in enumerate(items):
        row = table.rows[r]
        row.cells[0].text = ''
        row.cells[1].text = ''
        p0 = row.cells[0].paragraphs[0]
        p1 = row.cells[1].paragraphs[0]
        run = p0.add_run(key)
        _set_font(run, preset.font_heading, preset.size_body - 1, bold=True)
        run = p1.add_run(str(value))
        _set_font(run, preset.font_body, preset.size_body - 1)
        _set_cell_shading(row.cells[0], 'F5F5F5')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)


def add_classification_banner(doc, preset, classification):
    if not classification or classification == '公开':
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f'【{classification}】')
    _set_font(run, preset.font_heading, preset.size_body, bold=True,
              color=RGBColor(0xB0, 0x00, 0x00))


def add_version_history(doc, preset, version, date_str, author):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('版本历史')
    _set_font(run, preset.font_heading, preset.size_section, bold=True)

    rows = [
        ['版本', '日期', '作者', '修改说明'],
        [version or 'V1.0', date_str, author or '未知', '首次创建'],
    ]
    render_table(doc, preset, rows, has_header=True)


def add_approval_block(doc, preset, approval_list=None):
    items = approval_list if approval_list else [
        {'role': '编制', 'name': ''},
        {'role': '审核', 'name': ''},
        {'role': '批准', 'name': ''},
    ]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('审批记录')
    _set_font(run, preset.font_heading, preset.size_section, bold=True)

    today = datetime.date.today().strftime('%Y-%m-%d')
    rows = [['角色', '姓名', '日期', '签字']]
    for item in items:
        role = item.get('role', '')
        name = item.get('name') or '__________'
        date_str = item.get('date') or (today if role == '编制' else '')
        rows.append([role, name, date_str, '__________'])
    render_table(doc, preset, rows, has_header=True)


# ============================================================
# 六、对外入口
# ============================================================

def create_word_doc(output_path, title='', content='', doc_number=None,
                    version='V1.0', classification='内部', author=None,
                    company_name=None, logo_path=None, approval=None,
                    doc_format=None, use_odoo=True,
                    force_version_history=None, force_approval=None,
                    force_classification_banner=None,
                    force_doc_meta_table=None,
                    force_title_block=None,
                    force_toc=None):
    """生成企业 Word 文档。

    v7.3：每种 preset 自带 show_classification_banner / show_doc_meta_table /
    show_title_block 默认值，CLI 用 --with-* / --no-* 覆盖。
    v7.6：新增 force_toc / --with-toc / --no-toc 显式控制目录是否生成；
    默认行为：preset.table_of_contents=True 且 H1+H2 标题数 ≥ 4 才生成。"""
    info = doc_core.resolve_company_info(
        overrides={'company_name': company_name, 'logo_path': logo_path},
        use_odoo=use_odoo,
    )
    missing = doc_core.company_info_missing(info)
    if missing:
        raise RuntimeError(doc_core.company_info_error_payload(missing))

    company = info['company_name']
    logo = info['logo_path']

    if not doc_format or doc_format == 'auto':
        doc_format = doc_core.detect_format(title, content)
    preset = doc_core.get_preset(doc_format)

    print(f'📄 使用文档规范: {preset.name}')
    print(f'🏢 公司: {company}')
    print(f'🖼  LOGO: {logo}')
    if not HAS_PYGMENTS:
        print('💡 提示：装 pygments 可启用代码块语法高亮 — pip install pygments')

    _reset_bookmarks()
    doc = Document()

    # 1) 文档核心属性（OA 检索 / 投标书系统索引）
    set_doc_core_properties(
        doc, title=title, author=author or company,
        subject=preset.name, category=preset.name,
        keywords=f'{preset.name},{company}',
        comments=f'由火一五文档技能 v{__doc__.split("v")[1].split()[0] if "v" in (__doc__ or "") else "7"} 生成',
    )

    for sec in doc.sections:
        sec.top_margin = Cm(preset.margin_top)
        sec.bottom_margin = Cm(preset.margin_bottom)
        sec.left_margin = Cm(preset.margin_left)
        sec.right_margin = Cm(preset.margin_right)
        sec.header_distance = Cm(1.5)
        sec.footer_distance = Cm(1.5)

    build_header(doc, preset, logo, company,
                 doc_number=doc_number,
                 classification=classification,
                 title=title)
    build_footer(doc, preset, company)

    normal_style = doc.styles['Normal']
    normal_style.font.name = preset.font_body
    normal_style.font.size = Pt(preset.size_body)

    want_banner = (force_classification_banner
                   if force_classification_banner is not None
                   else getattr(preset, 'show_classification_banner', True))
    want_meta = (force_doc_meta_table
                 if force_doc_meta_table is not None
                 else getattr(preset, 'show_doc_meta_table', True))
    want_title = (force_title_block
                  if force_title_block is not None
                  else getattr(preset, 'show_title_block', True))

    # H1 dedup：若 markdown 首个 H1 与 --title 同文，且 preset.dedupe_h1_title=True，
    # 跳过该 H1，避免"大标题"和"H1"叠在一起。
    content_to_render = _maybe_dedupe_h1_title(content, title, preset, want_title)

    # v7.5.2: 如果 markdown 正文顶部已有元数据块（LLM 通常会写一份），
    # 跳过 CLI 自动生成的 文档编号/版本/密级/日期 表，避免和正文里的元数据
    # 表叠加显示。CLI 显式 --with-doc-meta-table 仍可强制开启。
    if (want_meta and force_doc_meta_table is None
            and _body_has_top_metadata_block(content_to_render)):
        want_meta = False

    if want_banner:
        add_classification_banner(doc, preset, classification)
    if want_title:
        add_title(doc, preset, title)
    if want_meta:
        add_doc_meta(doc, preset, doc_number, version, classification, author)

    # 2) 自动 TOC：preset.table_of_contents=True 且文档有 ≥4 H1+H2 时生成；
    #    CLI --with-toc / --no-toc 显式覆盖。
    if force_toc is not None:
        want_toc = force_toc
    else:
        want_toc = preset.table_of_contents and _heading_count_ge(
            content_to_render, 4)
    _reset_toc_collector()
    toc_paragraph, toc_placeholder = (None, None)
    if want_toc:
        toc_paragraph, toc_placeholder = add_toc_field(
            doc, preset, levels='1-3', title='目录')
        _set_update_fields_on_open(doc)

    render_content(doc, preset, content_to_render)

    # 3) 静态 TOC 回填：把渲染过程中收集到的 H1-H3 标题写到 TOC 字段缓存。
    #    用户打开 Word/WPS 之前能看到完整目录列表（无页码）；
    #    打开后 updateFields=true 触发刷新，替换为带页码的真目录。
    if toc_paragraph is not None:
        _finalize_static_toc(toc_paragraph, toc_placeholder, preset)

    want_version = (force_version_history
                    if force_version_history is not None
                    else preset.has_version_history)
    if want_version:
        add_version_history(doc, preset, version,
                            datetime.date.today().strftime('%Y-%m-%d'),
                            author or '未知')

    want_approval = (force_approval if force_approval is not None
                     else preset.has_approval)
    if want_approval or approval:
        add_approval_block(doc, preset, approval)

    doc.save(output_path)
    print(f'✅ 文档已生成: {output_path}')
    return output_path


# ============================================================
# 七、CLI
# ============================================================

def _use_legacy_cli(argv):
    if len(argv) <= 1:
        return False
    return not argv[1].startswith('-')


def _parse_args(argv):
    parser = argparse.ArgumentParser(
        prog='create-word-doc',
        description='火一五企业级 Word 生成器 v7.0（12 类规范）',
    )
    parser.add_argument('--output', '-o', required=False, default=None,
                        help='输出 .docx 路径（--list-formats 模式下可省略）')
    parser.add_argument('--title', default='', help='文档标题')
    parser.add_argument('--content', default='',
                        help='正文（Markdown）；以 @file 开头时读取文件')
    parser.add_argument('--doc-number', default=None)
    parser.add_argument('--version', default='V1.0')
    parser.add_argument('--classification', default='内部',
                        help='密级：公开/内部/秘密')
    parser.add_argument('--author', default=None)
    parser.add_argument('--doc-format', default='auto',
                        choices=['auto'] + doc_core.list_format_names())
    parser.add_argument('--company-name', default=None)
    parser.add_argument('--logo-path', default=None)
    parser.add_argument('--no-odoo', action='store_true')
    parser.add_argument('--with-version-history', action='store_true')
    parser.add_argument('--no-version-history', action='store_true')
    parser.add_argument('--with-approval', action='store_true')
    parser.add_argument('--no-approval', action='store_true')
    # v7.3：文档壳开关（默认值由 preset 决定，CLI 可覆盖）
    parser.add_argument('--with-classification-banner', action='store_true',
                        help='顶部右上 【内部】红字 banner（默认随 preset）')
    parser.add_argument('--no-classification-banner', action='store_true',
                        help='隐藏【内部】banner（合同 / 演讲稿 / 手册等默认隐藏）')
    parser.add_argument('--with-doc-meta-table', action='store_true',
                        help='文档编号/版本/密级/日期 顶部表格（默认随 preset）')
    parser.add_argument('--no-doc-meta-table', action='store_true',
                        help='隐藏顶部元数据表（合同 / 商业计划书 / 手册默认隐藏）')
    parser.add_argument('--with-title-block', action='store_true',
                        help='把 --title 渲染为大标题块（默认开）')
    parser.add_argument('--no-title-block', action='store_true',
                        help='不另起标题块（已在 markdown 里写 # 标题时用）')
    parser.add_argument('--with-toc', action='store_true',
                        help='强制生成目录（即使章节较少）')
    parser.add_argument('--no-toc', action='store_true',
                        help='不生成目录（默认 preset 启用 TOC 但章节 < 4 时自动跳过）')
    parser.add_argument('--list-formats', action='store_true',
                        help='打印 17 种 preset 名称及说明后退出')
    return parser.parse_args(argv[1:])


def _load_content(content_arg):
    if content_arg and content_arg.startswith('@'):
        path = content_arg[1:]
        with open(path, 'r', encoding='utf-8') as fh:
            return fh.read()
    return content_arg


def _flag_tristate(on, off):
    if on and off:
        return None
    if on:
        return True
    if off:
        return False
    return None


def main(argv=None):
    argv = argv if argv is not None else sys.argv
    try:
        if _use_legacy_cli(argv):
            create_word_doc(
                output_path=argv[1] if len(argv) > 1 else 'output.docx',
                title=argv[2] if len(argv) > 2 else '',
                content=argv[3] if len(argv) > 3 else '',
                doc_number=argv[4] if len(argv) > 4 else None,
                version=argv[5] if len(argv) > 5 else 'V1.0',
                classification=argv[6] if len(argv) > 6 else '内部',
                doc_format=argv[7] if len(argv) > 7 else 'auto',
            )
        else:
            args = _parse_args(argv)
            if args.list_formats:
                for name in doc_core.list_format_names():
                    p = doc_core.get_preset(name)
                    desc = getattr(p, 'description', '') or ''
                    flags = []
                    if getattr(p, 'show_classification_banner', True):
                        flags.append('banner')
                    if getattr(p, 'show_doc_meta_table', True):
                        flags.append('meta')
                    if p.has_version_history:
                        flags.append('版本史')
                    if p.has_approval:
                        flags.append('审批')
                    if p.table_of_contents:
                        flags.append('TOC')
                    flag_str = ' / '.join(flags) if flags else '简洁'
                    print(f'  {name:<8s} [{flag_str}]')
                    if desc:
                        print(f'           {desc}')
                return 0
            if not args.output:
                print('错误：--output 必填（除非用 --list-formats）',
                      file=sys.stderr)
                return 2
            create_word_doc(
                output_path=args.output,
                title=args.title,
                content=_load_content(args.content),
                doc_number=args.doc_number,
                version=args.version,
                classification=args.classification,
                author=args.author,
                company_name=args.company_name,
                logo_path=args.logo_path,
                doc_format=args.doc_format,
                use_odoo=not args.no_odoo,
                force_version_history=_flag_tristate(
                    args.with_version_history, args.no_version_history),
                force_approval=_flag_tristate(
                    args.with_approval, args.no_approval),
                force_classification_banner=_flag_tristate(
                    args.with_classification_banner,
                    args.no_classification_banner),
                force_doc_meta_table=_flag_tristate(
                    args.with_doc_meta_table, args.no_doc_meta_table),
                force_title_block=_flag_tristate(
                    args.with_title_block, args.no_title_block),
                force_toc=_flag_tristate(args.with_toc, args.no_toc),
            )
    except RuntimeError as exc:
        print(str(exc), file=sys.stderr)
        return 2
    return 0


if __name__ == '__main__':
    sys.exit(main())
